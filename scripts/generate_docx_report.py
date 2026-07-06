#!/usr/bin/env python3
"""Generate a formal DOCX QA report from a Markdown report.

Follows the `docx-test-report` skill standards: Calibri body, severity-shaded
bug rows, screenshots embedded inline at 14cm with italic captions.

Usage:
    python3 generate_docx_report.py runs/<date>-<target>/report.md
    python3 generate_docx_report.py runs/<date>-<target>/report.md -o reports/MyApp-QA-Report-20260627.docx

The Markdown is converted pragmatically:
- `#`/`##`/`###`      -> Heading 1/2/3
- `| a | b |` tables  -> Word tables (rows whose first cell is a severity
                         S1/S2/S3/S4 are colour-shaded)
- ```code```          -> Courier New 8pt block
- `- ` / `* `         -> bullet list
- `![caption](path)`  -> embedded image (path relative to the .md file) + caption
- everything else     -> body paragraph (with **bold** / `code` inline runs)
"""
from __future__ import annotations

import argparse
import os
import re
import sys
from datetime import date

try:
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("python-docx is required. Install with: pip install python-docx")

SEVERITY_COLOURS = {
    "S1": "FF4444", "Critical": "FF4444",
    "S2": "FF8C00", "High": "FF8C00", "Major": "FF8C00",
    "S3": "FFD700", "Medium": "FFD700", "Minor": "FFD700",
    "S4": "ADD8E6", "Low": "ADD8E6", "Cosmetic": "ADD8E6",
    "Info": "D3D3D3",
}

IMG_RE = re.compile(r"!\[(?P<cap>.*?)\]\((?P<path>.*?)\)")
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def shade_cell(cell, hex_colour: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_colour)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_inline_runs(paragraph, text: str) -> None:
    """Render **bold** and `code` inline spans inside a paragraph."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            paragraph.add_run(part)


def set_base_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def emit_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Light Grid Accent 1"
    header = rows[0]
    body = [r for r in rows[1:] if not all(set(c) <= set("-: ") for c in r)]
    cells = table.add_row().cells
    for i, txt in enumerate(header):
        if i < cols:
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(txt)
            run.bold = True
    for row in body:
        cells = table.add_row().cells
        sev = row[0] if row else ""
        colour = SEVERITY_COLOURS.get(sev)
        for i in range(cols):
            txt = row[i] if i < len(row) else ""
            cells[i].text = ""
            add_inline_runs(cells[i].paragraphs[0], txt)
            if colour:
                shade_cell(cells[i], colour)


def emit_image(doc: Document, base_dir: str, path: str, caption: str) -> None:
    abs_path = path if os.path.isabs(path) else os.path.join(base_dir, path)
    if os.path.exists(abs_path):
        try:
            doc.add_picture(abs_path, width=Cm(14))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as exc:  # corrupt/unsupported image
            _evidence_note(doc, f"[image could not be embedded: {abs_path} — {exc}]")
            return
        if caption:
            cap = doc.add_paragraph(f"Figure: {caption}")
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(9)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        _evidence_note(doc, f"Evidence (missing file): {abs_path}")


def _evidence_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)


def emit_code(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["No Spacing"]
    except KeyError:
        pass
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(8)


def convert(md_path: str, out_path: str) -> None:
    base_dir = os.path.dirname(os.path.abspath(md_path))
    with open(md_path, encoding="utf-8") as fh:
        lines = fh.read().splitlines()

    doc = Document()
    set_base_style(doc)

    i = 0
    table_buf: list[list[str]] = []
    code_buf: list[str] | None = None
    bullets: list[str] = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            emit_table(doc, table_buf)
            table_buf = []

    def flush_bullets():
        nonlocal bullets
        for b in bullets:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, b)
        bullets = []

    while i < len(lines):
        line = lines[i]

        if code_buf is not None:
            if line.strip().startswith("```"):
                emit_code(doc, code_buf)
                code_buf = None
            else:
                code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_table(); flush_bullets()
            code_buf = []
            i += 1
            continue

        if line.strip().startswith("|") and line.strip().endswith("|"):
            flush_bullets()
            table_buf.append(parse_table_row(line))
            i += 1
            continue
        else:
            flush_table()

        img = IMG_RE.search(line)
        if img:
            flush_bullets()
            emit_image(doc, base_dir, img.group("path"), img.group("cap"))
            i += 1
            continue

        if line.startswith("### "):
            flush_bullets(); doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            flush_bullets(); doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            flush_bullets(); doc.add_heading(line[2:].strip(), level=1)
        elif line.lstrip().startswith(("- ", "* ")):
            bullets.append(line.lstrip()[2:])
        elif line.strip() == "":
            flush_bullets()
        elif set(line.strip()) <= set("-") and len(line.strip()) >= 3:
            pass  # horizontal rule
        else:
            flush_bullets()
            p = doc.add_paragraph()
            add_inline_runs(p, line)
        i += 1

    flush_table(); flush_bullets()
    if code_buf is not None:
        emit_code(doc, code_buf)

    os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")


def default_out(md_path: str) -> str:
    # Expected: projects/<slug>/runs/<YYYY-MM-DD>/report.md
    base_dir = os.path.dirname(os.path.abspath(md_path))   # .../runs/<date>
    runs_dir = os.path.dirname(base_dir)                    # .../runs
    if os.path.basename(runs_dir) == "runs":
        project_root = os.path.dirname(runs_dir)           # .../projects/<slug>
        name = os.path.basename(project_root)              # <slug>
    else:
        project_root = base_dir
        name = os.path.basename(base_dir) or "App"
    reports = os.path.join(project_root, "reports")
    stamp = date.today().strftime("%Y%m%d")
    return os.path.join(reports, f"{name}-QA-Report-{stamp}.docx")


def main() -> None:
    ap = argparse.ArgumentParser(description="Markdown QA report -> DOCX")
    ap.add_argument("markdown", help="path to the markdown report (e.g. runs/<date>/report.md)")
    ap.add_argument("-o", "--output", help="output .docx path")
    args = ap.parse_args()

    if not os.path.exists(args.markdown):
        sys.exit(f"Not found: {args.markdown}")
    out = args.output or default_out(args.markdown)
    convert(args.markdown, out)


if __name__ == "__main__":
    main()
